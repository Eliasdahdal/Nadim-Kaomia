import streamlit as st
import random
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import zipfile
from docx.shared import Inches, Cm


# Function to convert numbers to Arabic numerals
def convert_to_arabic_numerals(number):
    arabic_numbers = str(number).translate(str.maketrans('0123456789', '٠١٢٣٤٥٦٧٨٩'))
    return arabic_numbers

# Function to set text direction to RTL and justify alignment for paragraphs
def set_rtl_and_justify(paragraph):
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justify alignment
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

# Function to set the entire document to RTL
def set_document_rtl(doc):
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        sectPr.append(bidi)

# Function to set the table to RTL
def set_table_rtl(table):
    tblPr = table._tblPr
    tblPr.append(OxmlElement('w:bidiVisual'))
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT

# Function to set cell properties for RTL
def set_cell_rtl(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(OxmlElement('w:rtl'))
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justify alignment
        pPr = paragraph._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)


# Function to ensure a paragraph or table stays together on one page
def set_keep_together(paragraph_or_table):
    if hasattr(paragraph_or_table, "rows"):  # For tables
        for row in paragraph_or_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    pPr = paragraph._element.get_or_add_pPr()
                    keep_lines = OxmlElement("w:keepLines")
                    keep_next = OxmlElement("w:keepNext")
                    pPr.append(keep_lines)
                    pPr.append(keep_next)
    else:  # For paragraphs
        pPr = paragraph_or_table._element.get_or_add_pPr()
        keep_lines = OxmlElement("w:keepLines")
        keep_next = OxmlElement("w:keepNext")
        pPr.append(keep_lines)
        pPr.append(keep_next)


# Function to add a large image to the first page content
def add_image_to_first_page(doc, image_path):
    # Adjust the top margin for the first section
    section = doc.sections[0]
    section.top_margin = Cm(0.5)  # Reduce top margin to 1 cm

    # Add the image as the first element in the document
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(6.5))  # Adjust the width to fit the page
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# Function to create a Word document for a single model
# Function to create a Word document for a single model
def create_word_file(image_header, footer, questions, true_false_statements):
    doc = Document()
    set_document_rtl(doc)  # Set the entire document to RTL

    # Set default paragraph style to RTL and justify alignment
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    style.font.size = Pt(12)

    # Apply RTL to style
    pPr = style.element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

    # Add the image to the first page content only
    if image_header:
        add_image_to_first_page(doc, image_header)

    # Add True/False section
    if true_false_statements:
        tf_intro = doc.add_paragraph()
        run = tf_intro.add_run("أولا: ضع كلمة صح أو غلط أمام العبارات التالية")
        run.bold = True  # Make text bold
        run.font.size = Pt(14)  # Increase font size
        set_rtl_and_justify(tf_intro)

        for statement in true_false_statements:
            tf_paragraph = doc.add_paragraph(f"{statement} (       )")
            set_rtl_and_justify(tf_paragraph)

    # Add a separator
    separator = doc.add_paragraph("-" * 80)
    set_rtl_and_justify(separator)

    # Add multiple-choice questions
    mcq_intro = doc.add_paragraph()
    run = mcq_intro.add_run("ثانيا : اختر الإجابة الصحيحة مما يلي")
    run.bold = True  # Make text bold
    run.font.size = Pt(14)  # Increase font size
    set_rtl_and_justify(mcq_intro)

    for idx, q in enumerate(questions, 1):
        arabic_number = convert_to_arabic_numerals(idx)
        question_paragraph = doc.add_paragraph(f"{arabic_number}- {q['question']}")
        set_rtl_and_justify(question_paragraph)
        set_keep_together(question_paragraph)  # Ensure question stays on one page

        # Add options in table format
        table = doc.add_table(rows=2, cols=2)
        set_table_rtl(table)
        table.style = 'Table Grid'

        options_labels = ['أ', 'ب', 'ج', 'د']
        for i, option in enumerate(q['options']):
            row_index = i // 2
            col_index = i % 2
            cell = table.cell(row_index, col_index)
            set_cell_rtl(cell)
            cell.text = f"{options_labels[i]}- {option}"

        # Ensure table stays on one page
        set_keep_together(table)

        spacer = doc.add_paragraph()
        set_rtl_and_justify(spacer)

    if footer:
        footer_paragraph = doc.add_paragraph(footer, style='Heading 2')
        set_rtl_and_justify(footer_paragraph)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit app
st.title("إنشاء نماذج أسئلة")

# Input: Number of headers
num_headers = st.number_input("عدد الترويسات:", min_value=1, value=1, step=1)

# Create a list to store header images and names
headers = []

# Allow the user to upload images and provide names for each header
for i in range(int(num_headers)):
    st.subheader(f"الترويسة {i + 1}")
    header_name = st.text_input(f"اسم الترويسة {i + 1}:", f"Header_{i + 1}").strip()
    header_image = st.file_uploader(f"تحميل صورة الترويسة {i + 1}:", type=["png", "jpg", "jpeg"], key=f"header_{i}")
    if header_image:
        headers.append({"name": header_name, "image": header_image})

# Footer input
footer = st.text_area("أدخل التذييل (Footer):", "انتهت الأسئلة").strip()

# True/False statements
st.subheader("إضافة عبارات صح وخطأ")
true_false_statements_input = st.text_area(
    "أدخل عبارات صح وخطأ (افصل كل عبارة بسطر جديد):",
    "ألا يعاقب المسؤول أمام مرؤوسيه من مهارات القيادة\n"
    "يقتضي القيام بعملية التحليل السياسي لتحديد كيفية التعامل مع الظاهرة\n"
    "المؤامرة في التحليل السياسي شرط لازم وكاف\n"
    "التحليل السياسي ترف فكري يمارسه البعض بقصد الدعاية"
)

# Clean up true/false statements
true_false_statements = [s.strip() for s in true_false_statements_input.split('\n') if s.strip()]

# Multiple choice questions
st.subheader("إضافة أسئلة اختيار من متعدد")
num_questions = st.number_input("عدد الأسئلة:", min_value=1, value=5, step=1)

questions = []
for i in range(num_questions):
    with st.expander(f"السؤال {i + 1}"):
        # Clean up question text
        question = st.text_input(f"السؤال {i + 1}:", f"سؤال {i + 1}").strip()
        options = []
        col1, col2 = st.columns(2)
        with col1:
            # Clean up options
            options.append(st.text_input(f"الخيار الأول للسؤال {i + 1}:", f"خيار 1").strip())
            options.append(st.text_input(f"الخيار الثالث للسؤال {i + 1}:", f"خيار 3").strip())
        with col2:
            options.append(st.text_input(f"الخيار الثاني للسؤال {i + 1}:", f"خيار 2").strip())
            options.append(st.text_input(f"الخيار الرابع للسؤال {i + 1}:", f"خيار 4").strip())
        questions.append({"question": question, "options": options})

# Number of models
num_models = st.number_input("عدد النماذج لكل ترويسة:", min_value=1, value=3, step=1)
file_name = st.text_input("اسم الملف (بدون الامتداد):", "quiz_model").strip()

# Generate models and create ZIP file
if st.button("إنشاء النماذج"):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zip_file:
        for header in headers:
            header_name = header["name"]
            header_image = header["image"]

            # Create models for the current header
            for i in range(1, num_models + 1):
                shuffled_questions = random.sample(questions, len(questions))
                shuffled_statements = random.sample(true_false_statements, len(true_false_statements))
                buffer = create_word_file(header_image, footer, shuffled_questions, shuffled_statements)

                # Add the file to the zip in a folder for the current header
                folder_name = header_name
                file_path = f"{folder_name}/{file_name}_{i}.docx"
                zip_file.writestr(file_path, buffer.getvalue())

    zip_buffer.seek(0)
    st.download_button(
        label="تحميل جميع النماذج كملف مضغوط",
        data=zip_buffer,
        file_name=f"{file_name}_models.zip",
        mime="application/zip"
    )
